from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pandas as pd
from mswordtree.Item import Item

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
            
def ParseTableToDataFrame(table):
    data = []
    keys = None
    
    gtext = [] # horizontal row text
    isVertical = True
    for i, row in enumerate(table.rows):
        text = list((cell.text for cell in row.cells))
        if i == 0:
            # if its a vertical table, its heading will have bold property
            try:
                #if all([cell.paragraphs[0].runs[0].font.bold for cell in row.cells]) or len(table.columns) > 2:

                isVertical = all([cell.paragraphs[0].runs[0].font.bold for cell in row.cells])
            except:
                isVertical = True
                
            if isVertical == True:            
                keys = tuple(text)                
                continue
            # else its a horizontal table
            else:  
                keys = [key for key in ((cell.text for cell in table.columns[0].cells))]                
                    
        if isVertical == False: 
            if len(text) > 0:            
                try:
                    va = text[1]
                    gtext.append(va)
                except:
                    continue
        else: 
            row_data = dict(zip(keys, text)) 
            data.append(row_data)
     
    if isVertical == False:        
        row_data = dict(zip(keys, gtext)) 
        data.append(row_data) 
        
    df = pd.DataFrame(data)
    df = df.fillna('')
    return df

def GetHeadingLevel(stylename):
    name = stylename
    
    try:
        
        level = [int(s) for s in name.split() if s.isdigit()][0]
    
        return level
    except:
        return 1

   

root = None
parent = root

def CreateHeading(block):
    global parent
    item = Item()
    item.Content = block.text.strip()
    item.Type = block.style.name   
    
    parentLevel = GetHeadingLevel(parent.Type) if parent.Type != 'root' else 0   
    itemLevel = GetHeadingLevel(item.Type)
    
    

    if parent.Parent:
        if (parent.Parent.Type == 'root' and parentLevel == itemLevel):
            root.Items.append(item) 
            item.Parent = root#item
            parent = item
            return item

    if parent.Type == 'root': 
        root.Items.append(item) 
        item.Parent = root#item
        parent = item
        return item

    # If the heading is new then add it to the root level
    if ((GetHeadingLevel(item.Type) == 1)): 
        root.Items.append(item) 
        item.Parent = root#item
        parent = item

    

    
    # Check if the new heading is a subheading of its parent or not
    elif ((parentLevel + 1) == itemLevel):
        parent.Items.append(item)
        item.Parent = parent
    
    # Get the last heading from the parent and then make that heading the parent of this newly created heading
    elif ((parentLevel + 1) < itemLevel):
        parent = GetLastHeading(parent)                          
        parent.Items.append(item)
        item.Parent = parent

    
    
        
    # Find the relivent parent and then add the new heading to its parent heading
    elif ((parentLevel + 1) > itemLevel): 
        parent = FindParent(parent, item)
        parent.Items.append(item)
        item.Parent = parent

        
    return item
    

def FindParent(par, item):
    parentLevel = GetHeadingLevel(par.Type)
    itemLevel = GetHeadingLevel(item.Type)    
   
    if((parentLevel + 1) > itemLevel):
        return FindParent(par.Parent, item)
    else:
        return par

    
def GetLastHeading(pare):
    for i in reversed(parent.Items):
            if 'Heading' in i.Type:
                return i
    return pare
                
def AddParagraph(head, block):
    item = Item()
    item.Type = block.style.name
    item.Content = block.text.strip()
    item.Parent = head
    head.Items.append(item)   
   
    pass

def AddTable(head, block):
    item = Item()
    item.Type = 'Table'
    item.Parent = head    
    item.Content = ParseTableToDataFrame(block)
    head.Items.append(item)
    pass

def GetWordDocTree(filename):
    global root
    global parent
    
    doc = Document(filename)
    root = Item()
    root.Type = "root"
    root.Content = filename
    parent = root
    
    head = root    
    for block in iter_block_items(doc):

        if ('Heading' in block.style.name) and len(block.text.strip()) > 0: #Heading
            
            head = CreateHeading(block)            
            pass
            
        
        if isinstance(block, Table): #Table             
            AddTable(head, block)
            pass
            
        
        elif len(block.text.strip()) > 0: # Paragraph and everything
            AddParagraph(head, block)
            pass
            
    root.Content = filename       
    return root        
